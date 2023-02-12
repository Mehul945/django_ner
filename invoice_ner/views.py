from django.shortcuts import render, HttpResponse, redirect
from .ner_extract import *
from PIL import Image
from django.contrib.auth.decorators import login_required
from django.contrib import auth, messages
from django.contrib.auth.models import User
import os
# Create your views here.
from django.http import JsonResponse
def custom_404(request, exception):
    return render(request, '404.html')
def get_text_fr_img(file_path):
    import docx
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

def docx_to_pdf(file_path,filename):
    import aspose.words as aw
    doc = aw.Document(file_path)
    doc.save(filename)
    os.remove(file_path)
def get_text_fr_pdf(file_path):
    from PyPDF2 import PdfReader
    reader = PdfReader(file_path)
    page = reader.pages[0]
    text = page.extract_text()
    return text


def home(request):
    return render(request, 'index.html')


def login(request):
    if request.user.is_authenticated:
        return redirect("/")
    elif request.method == "POST":
        email = request.POST.get('email')
        password = request.POST.get('password')
        user = auth.authenticate(username=email.split(
            '@')[0], email=email, password=password)
        auth.login(request,user)
        messages.error(request, 'Login Success')
        return redirect("/")
    return render(request, 'login.html')

def logout(request):
    auth.logout(request)
    return redirect('/')

def signup(request):
    if request.method == "POST":
        email = request.POST['email']
        password = request.POST['password']
        password2 = request.POST['verify-password']
        if password != password2:
            messages.info(request, 'password didn"t match')
            return redirect('signup')
        if User.objects.filter(email=email).exists():
            messages.info(request, 'Email Taken')
            return redirect('signup')
        else:
            user = User.objects.create_user(
                email=email, username=email.split('@')[0], password=password)
            user.save()
            return redirect('login')
    return render(request, 'signup.html')


@login_required
def invoice_extract(request):
    if request.method == "POST":
        file = request.FILES.get("doc_file")
        type = request.POST.get('type')
        text = ''
        if type=='docx':
            with open(f'static/{request.user.username}.docx','wb') as fh:
                for chunk in file.chunks():
                    fh.write(chunk)
            docx_to_pdf(f'static/{request.user.username}.docx',f'static/{request.user.username}.pdf')
            text=get_text_fr_pdf(f'static/{request.user.username}.pdf')
            print(text)
        elif type=='PDF':
            with open(f'static/{request.user.username}.pdf','wb') as fh:
                for chunk in file.chunks():
                    fh.write(chunk)
            text=get_text_fr_pdf(f'static/{request.user.username}.pdf')
        else:
            with open(f'static/{request.user.username}_{file.name}','wb') as fh:
                for chunk in file.chunks():
                    fh.write(chunk)
            text=extract_text_from_image(file,request.user.username)
        doc = convert_doc(text)
        sents = [i.text for i in doc.sents]
        bank_details = []
        if type == 'bank passbook':
            bank_details = extract_bank(sents)
        extract_entities(doc,request.user.username)
        contacts = extract_contacts(doc)
        data = json.load(open(fr'result_files/{request.user.username}.json'))
        return render(request, 'result.html', {"data": data})
    return redirect("/")


@login_required
def update_data(request):
    if request.method == "POST":
        data = json.load(
            open(fr'/home/mehulspec4513/NER_project/result_files/{request.user.username}.json'))
        value = request.POST.get("value")
        key = request.POST.get("key")
        index = request.POST.get("index")
        data[key][int(index)] = value
        print(data)
        json.dump(data, open(
            rf'result_files/{request.user.username}.json', 'w+'))
    return HttpResponse(f"{value}")


@login_required
def download_json(request):
    file_path = f'result_files/{request.user.username}.json'
    with open(file_path, 'r') as fh:
        response = HttpResponse(fh.read(), content_type="application/json")
        response['Content-Disposition'] = 'attachment; filename=' + request.user.username
        return response
    return HttpResponse("")


@login_required
def download_txt(request):
    file_path = f'result_files/{request.user.username}.txt'
    with open(file_path, 'r') as fh:
        response = HttpResponse(fh.read(), content_type="application/txt")
        response['Content-Disposition'] = 'attachment; filename=' + request.user.username
        return response
    return HttpResponse("")
